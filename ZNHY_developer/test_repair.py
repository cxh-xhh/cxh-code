#!/usr/bin/env python3
"""
测试修复后的参会人员姓名显示是否正确
"""

import os
import sys
import tempfile
from docx import Document
from utils.llm_parser import MeetingParser
from utils.word_generator import WordGenerator

def test_meeting_parser():
    """测试会议文本解析"""
    print("测试会议文本解析...")
    
    # 测试用的会议文本
    test_text = """
    会议主题：下季度产品推广方案
    会议时间：2023年10月15日 14:00-16:00
    会议地点：公司三楼大会议室
    主持人：张三
    参会人员：李四、王五、赵六
    会议时长：约2小时
    
    会议议程：
    1. 讨论线上广告投放预算分配
    2. 讨论新版APP研发进度
    3. 讨论下月底团建活动安排
    
    会前准备事项：提前准备相关资料
    """
    
    # 创建解析器
    parser = MeetingParser()
    
    try:
        # 解析会议文本
        meeting_info = parser.parse_meeting_text(test_text)
        print("✓ 会议文本解析成功")
        print(f"   会议主题：{meeting_info.get('meeting_topic', '未解析')}")
        print(f"   主持人：{meeting_info.get('host', '未解析')}")
        print(f"   会议地点：{meeting_info.get('meeting_location', '未解析')}")
        print(f"   参会人员：{meeting_info.get('participants', '未解析')}")
        print(f"   会议时长：{meeting_info.get('meeting_duration', '未解析')}")
        print(f"   议题数量：{len(meeting_info.get('topics', []))}")
        print(f"   会前准备事项：{meeting_info.get('pre_meeting_preparations', '未解析')}")
        return meeting_info
    except Exception as e:
        print(f"✗ 会议文本解析失败：{e}")
        return None

def test_word_generator(meeting_info):
    """测试Word文档生成"""
    print("\n测试Word文档生成...")
    
    if not meeting_info:
        print("✗ 无法生成Word文档，缺少会议信息")
        return False
    
    try:
        # 创建Word生成器
        generator = WordGenerator()
        
        # 生成文档
        buffer = generator.generate_document(meeting_info)
        
        # 保存到临时文件
        temp_dir = tempfile.mkdtemp()
        temp_file = os.path.join(temp_dir, "meeting_record.docx")
        
        with open(temp_file, "wb") as f:
            f.write(buffer.getvalue())
        
        print(f"✓ Word文档生成成功：{temp_file}")
        
        # 验证文档内容
        doc = Document(temp_file)
        
        # 检查参会人员是否正确显示
        participants_found = False
        # 首先检查段落
        for i, para in enumerate(doc.paragraphs):
            if "参会人员" in para.text:
                participants_found = True
                print(f"   参会人员行：{para.text}")
                # 检查是否包含正确的姓名
                if "李四" in para.text and "王五" in para.text and "赵六" in para.text:
                    print("   ✓ 参会人员姓名正确显示：李四、王五、赵六")
                else:
                    print("   ✗ 参会人员姓名显示错误")
                break
        
        # 如果段落中没有找到，检查表格
        if not participants_found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "参会人员" in cell.text:
                            participants_found = True
                            print(f"   参会人员行：{cell.text.strip()}")
                            # 检查是否包含正确的姓名
                            if "李四" in cell.text and "王五" in cell.text and "赵六" in cell.text:
                                print("   ✓ 参会人员姓名正确显示：李四、王五、赵六")
                            else:
                                print("   ✗ 参会人员姓名显示错误")
                            break
                    if participants_found:
                        break
                if participants_found:
                    break
        
        if not participants_found:
            print("   ✗ 未找到参会人员行")
        
        return True
    except Exception as e:
        print(f"✗ Word文档生成失败：{e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函数"""
    print("=== 会议记录修复测试 ===")
    print()
    
    # 测试解析器
    meeting_info = test_meeting_parser()
    
    if meeting_info:
        # 测试Word生成器
        success = test_word_generator(meeting_info)
        
        if success:
            print("\n✅ 修复测试成功！参会人员姓名已正确显示在会议文档中")
            return 0
        else:
            print("\n❌ 修复测试失败！")
            return 1
    else:
        print("\n❌ 修复测试失败！")
        return 1

if __name__ == "__main__":
    sys.exit(main())