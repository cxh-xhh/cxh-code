from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_meeting_record(meeting_info):
    """
    创建会议记录 Word 文档
    :param meeting_info: 包含会议信息的字典
    :return: 生成的 Word 文档对象
    """
    # 创建文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 1. 标题：会议记录
    title = doc.add_paragraph()
    title_run = title.add_run('会议记录')
    title_run.font.size = Pt(18)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # 空行

    # 2. 会议主题
    doc.add_paragraph('会议主题', style='Heading 1')
    doc.add_paragraph(meeting_info.get('meeting_topic', ''))
    doc.add_paragraph()

    # 3. 主持人
    doc.add_paragraph('主持人', style='Heading 1')
    doc.add_paragraph(meeting_info.get('host', ''))
    doc.add_paragraph()

    # 4. 会议地点
    doc.add_paragraph('会议地点', style='Heading 1')
    doc.add_paragraph(meeting_info.get('meeting_location', ''))
    doc.add_paragraph()

    # 5. 参会人员
    doc.add_paragraph('参会人员', style='Heading 1')
    doc.add_paragraph(meeting_info.get('participants', ''))
    doc.add_paragraph()

    # 6. 会议时长
    doc.add_paragraph('会议时长', style='Heading 1')
    doc.add_paragraph(meeting_info.get('meeting_duration', ''))
    doc.add_paragraph()

    # 7. 会议内容记录
    doc.add_paragraph('会议内容记录', style='Heading 1')
    doc.add_paragraph()

    # 处理议题
    topics = meeting_info.get('topics', [])
    for i, topic in enumerate(topics, 1):
        # 议题标题
        topic_title = doc.add_paragraph()
        topic_title_run = topic_title.add_run(f'议题{i}：{topic.get("title", "")}')
        topic_title_run.bold = True

        # 负责人
        if topic.get('leader'):
            doc.add_paragraph(f'负责人：{topic.get("leader")}')

        # 会前准备
        if topic.get('preparation'):
            doc.add_paragraph(f'会前准备：{topic.get("preparation")}')

        # 参与人员
        if topic.get('participants'):
            doc.add_paragraph(f'参与人员：{topic.get("participants")}')

        doc.add_paragraph()  # 空行

    # 8. 会前准备事项
    if meeting_info.get('pre_meeting_preparations'):
        pre_meeting = doc.add_paragraph()
        pre_meeting_run = pre_meeting.add_run('会前准备事项：')
        pre_meeting_run.bold = True
        pre_meeting.add_run(f' {meeting_info.get("pre_meeting_preparations")}')

    return doc


if __name__ == '__main__':
    # 测试数据
    test_info = {
        'meeting_topic': '下季度产品推广方案',
        'meeting_location': '公司三楼大会议室',
        'participants': '李明、张娜、王磊、赵晓雨',
        'meeting_duration': '约两小时',
        'topics': [
            {
                'title': '下季度产品推广方案（线上广告投放预算分配）',
                'leader': '李明',
                'preparation': '准备相关数据'
            },
            {
                'title': '新版 APP 研发进度',
                'leader': '王磊'
            },
            {
                'title': '下月底团建活动安排（备选地点：近郊民宿、爬山）',
                'leader': '张娜',
                'participants': '全体参会者（赵晓雨参与讨论）'
            }
        ],
        'pre_meeting_preparations': '提前将会议资料发到群里'
    }

    doc = create_meeting_record(test_info)
    doc.save('测试会议记录.docx')
    print('测试会议记录已生成：测试会议记录.docx')