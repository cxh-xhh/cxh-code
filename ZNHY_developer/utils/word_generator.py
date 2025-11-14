from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import io
from typing import Dict, Any

class WordGenerator:
    def __init__(self):
        self.document = Document()
    
    def generate_document(self, meeting_data: Dict[str, Any]) -> io.BytesIO:
        """生成Word文档"""
        self._setup_document()
        self._add_title()
        self._add_meeting_info_table(meeting_data)
        self._add_meeting_content(meeting_data)
        
        return self._save_to_buffer()
    
    def _setup_document(self):
        """设置文档基本样式"""
        # 设置中文字体
        self.document.styles['Normal'].font.name = '微软雅黑'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _add_title(self):
        """添加标题"""
        title = self.document.add_heading('会议记录', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 标题样式
        for run in title.runs:
            run.font.size = Pt(16)
            run.font.bold = True
    
    def _add_meeting_info_table(self, meeting_data: Dict[str, Any]):
        """添加会议信息表格"""
        # 创建4列5行的表格
        table = self.document.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        
        # 设置表格宽度
        for row in table.rows:
            row.height = Inches(0.3)
        
        # 填充表头和数据
        cells = table.rows[0].cells
        cells[0].text = '会议主题'
        cells[1].text = meeting_data.get('meeting_topic', '')
        cells[2].text = '主持人'
        cells[3].text = meeting_data.get('host', '')
        
        # 合并第二行单元格
        self._merge_cells(table.rows[1].cells, 0, 3)
        table.rows[1].cells[0].text = f"会议地点：{meeting_data.get('meeting_location', '')}"
        
        # 合并第三行单元格
        self._merge_cells(table.rows[2].cells, 0, 3)
        participants = meeting_data.get('participants', [])
        # 处理参会人员格式，支持字符串列表或字典列表
        if participants and isinstance(participants[0], dict):
            participant_names = [p.get('name', '') for p in participants]
        else:
            participant_names = participants
        participants_text = '、'.join(participant_names)
        table.rows[2].cells[0].text = f"参会人员：{participants_text}"
        
        # 合并第四行单元格
        self._merge_cells(table.rows[3].cells, 0, 3)
        table.rows[3].cells[0].text = f"会议时长：{meeting_data.get('meeting_duration', '')}"
        
        # 合并第五行单元格
        self._merge_cells(table.rows[4].cells, 0, 3)
        table.rows[4].cells[0].text = "会议内容记录"
        
        # 设置单元格样式
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
    
    def _merge_cells(self, cells, start_idx, end_idx):
        """合并单元格"""
        if start_idx < end_idx:
            cells[start_idx].merge(cells[end_idx])
    
    def _add_meeting_content(self, meeting_data: Dict[str, Any]):
        """添加会议内容"""
        self.document.add_paragraph()  # 空行
        
        # 处理议题内容（新格式）
        topics = meeting_data.get('topics', [])
        if topics:
            for i, topic in enumerate(topics, 1):
                # 添加议题标题
                topic_title = self.document.add_paragraph()
                topic_title.add_run(f'议题{i}：{topic.get("title", "")}').bold = True
                
                # 添加负责人
                if topic.get('leader'):
                    responsible_para = self.document.add_paragraph()
                    responsible_para.add_run(f'负责人：{topic["leader"]}')
                
                # 添加会前准备
                if topic.get('preparation'):
                    prep_para = self.document.add_paragraph()
                    prep_para.add_run(f'会前准备：{topic["preparation"]}')
                
                self.document.add_paragraph()  # 空行
        
        # 处理旧格式的agenda_items（用于向后兼容）
        elif 'agenda_items' in meeting_data:
            agenda_items = meeting_data.get('agenda_items', [])
            for i, item in enumerate(agenda_items, 1):
                topic_title = self.document.add_paragraph()
                topic_title.add_run(f'议题{i}：{item.get("topic", "")}').bold = True
                
                if item.get('responsible_person'):
                    responsible_para = self.document.add_paragraph()
                    responsible_para.add_run(f'负责人：{item["responsible_person"]}')
                
                if item.get('preparation'):
                    prep_para = self.document.add_paragraph()
                    prep_para.add_run(f'会前准备：{item["preparation"]}')
                
                self.document.add_paragraph()  # 空行
        
        # 添加会前准备事项
        pre_meeting_preparations = meeting_data.get('pre_meeting_preparations', '')
        if pre_meeting_preparations:
            prep_para = self.document.add_paragraph()
            prep_para.add_run(f'会前准备事项：{pre_meeting_preparations}').bold = True
    
    def _save_to_buffer(self) -> io.BytesIO:
        """保存到内存缓冲区"""
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer